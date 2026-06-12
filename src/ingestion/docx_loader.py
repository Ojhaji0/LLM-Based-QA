"""
DOCX Loader — Extracts text from Word documents.
"""
from loguru import logger
import docx
import os

def load_docx(filepath: str) -> list[dict]:
    """
    Extract text from a DOCX file and return it as a list of block dictionaries.
    Simulates the structure returned by pdf_loader for compatibility.
    """
    logger.info(f"Loading DOCX: {filepath}")
    blocks = []
    
    try:
        doc = docx.Document(filepath)
        filename = os.path.basename(filepath)
        
        # Extract paragraphs
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                blocks.append({
                    "text": text,
                    "page": 1,  # DOCX doesn't have explicit pages easily, default to 1
                    "is_table": False,
                    "source_file": filename,
                    "extraction_method": "python-docx"
                })
                
        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    table_text.append(row_text)
            
            if table_text:
                blocks.append({
                    "text": "\n".join(table_text),
                    "page": 1,
                    "is_table": True,
                    "source_file": filename,
                    "extraction_method": "python-docx"
                })
                
    except Exception as e:
        logger.error(f"Error loading {filepath}: {e}")
        raise
        
    return blocks
